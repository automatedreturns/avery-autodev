"""
DOCX Processor - Extract and structure content from Word documents
Handles: .docx files (Microsoft Word)
"""
import logging
import os
from typing import Dict, Any, Optional, List
import time

logger = logging.getLogger(__name__)


class DOCXProcessor:
    """Process DOCX documents with content extraction and AI structuring"""

    def __init__(self, ai_service):
        """
        Initialize DOCX processor.

        Args:
            ai_service: AIModelService instance for content analysis
        """
        self.ai_service = ai_service

    async def process(self, file_path: str) -> Dict[str, Any]:
        """
        Process DOCX file with content extraction and AI analysis.

        Args:
            file_path: Path to DOCX file

        Returns:
            {
                "filename": str,
                "type": str,
                "size_bytes": int,
                "paragraph_count": int,
                "processed": bool,
                "ai_provider": str,
                "ai_model": str,
                "processing_time_ms": int,
                "extracted_content": {
                    "document_type": str,
                    "summary": str,
                    "main_points": list[str],
                    "action_items": list[str],
                    "test_scenarios": list[dict],
                    "tables": list[dict],
                    "raw_text": str (truncated)
                },
                "error": Optional[str]
            }
        """
        start_time = time.time()

        result = {
            "filename": os.path.basename(file_path),
            "type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "size_bytes": os.path.getsize(file_path),
            "paragraph_count": 0,
            "processed": False,
            "ai_provider": "none",
            "ai_model": "none",
            "processing_time_ms": 0,
            "extracted_content": {
                "document_type": "",
                "summary": "",
                "main_points": [],
                "action_items": [],
                "test_scenarios": [],
                "tables": [],
                "raw_text": ""
            },
            "error": None
        }

        try:
            # Extract content from DOCX
            logger.info(f"Extracting content from DOCX: {file_path}")
            extraction_result = await self._extract_content(file_path)

            if not extraction_result["success"]:
                result["error"] = extraction_result["error"]
                return result

            result["paragraph_count"] = extraction_result["paragraph_count"]
            raw_text = extraction_result["text"]
            result["extracted_content"]["tables"] = extraction_result.get("tables", [])

            # Store truncated raw text
            result["extracted_content"]["raw_text"] = raw_text[:2000] + ("..." if len(raw_text) > 2000 else "")

            # Use AI to analyze and structure the content
            logger.info(f"Analyzing DOCX content with AI")
            analysis_result = await self._analyze_with_ai(raw_text, result["paragraph_count"])

            if analysis_result["success"]:
                result["processed"] = True
                result["ai_provider"] = analysis_result["provider"]
                result["ai_model"] = analysis_result["model"]

                # Parse AI response
                parsed_content = self._parse_analysis_response(analysis_result["content"])
                result["extracted_content"].update(parsed_content)

                logger.info(f"Successfully processed DOCX with {analysis_result['provider']}")
            else:
                # If AI fails, at least provide raw text
                result["processed"] = True
                result["error"] = f"AI analysis failed: {analysis_result.get('error')}, providing raw text only"
                result["extracted_content"]["summary"] = f"Word document with {result['paragraph_count']} paragraphs. AI analysis unavailable."

        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}", exc_info=True)
            result["error"] = str(e)

        # Calculate processing time
        result["processing_time_ms"] = int((time.time() - start_time) * 1000)

        return result

    async def _extract_content(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text, tables, and formatting from DOCX.

        Args:
            file_path: Path to DOCX file

        Returns:
            {
                "success": bool,
                "text": str,
                "paragraph_count": int,
                "tables": list[dict],
                "error": Optional[str]
            }
        """
        try:
            from docx import Document

            doc = Document(file_path)

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Preserve some formatting info
                    style = para.style.name if para.style else "Normal"
                    if style.startswith("Heading"):
                        paragraphs.append(f"\n## {text}")
                    elif style == "List Bullet" or style == "List Number":
                        paragraphs.append(f"  • {text}")
                    else:
                        paragraphs.append(text)

            combined_text = "\n".join(paragraphs)
            paragraph_count = len([p for p in paragraphs if p.strip()])

            # Extract tables
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                if table_data:
                    tables.append({
                        "table_index": table_idx + 1,
                        "rows": len(table_data),
                        "columns": len(table_data[0]) if table_data else 0,
                        "data": table_data[:10]  # Store first 10 rows only
                    })

            return {
                "success": True,
                "text": combined_text,
                "paragraph_count": paragraph_count,
                "tables": tables,
                "error": None
            }

        except ImportError:
            logger.error("python-docx not installed")
            return {
                "success": False,
                "text": "",
                "paragraph_count": 0,
                "tables": [],
                "error": "python-docx library not installed"
            }

        except Exception as e:
            logger.error(f"DOCX content extraction failed: {e}")
            return {
                "success": False,
                "text": "",
                "paragraph_count": 0,
                "tables": [],
                "error": str(e)
            }

    async def _analyze_with_ai(self, text: str, paragraph_count: int) -> Dict[str, Any]:
        """
        Analyze DOCX content with AI to extract structured information.

        Args:
            text: Extracted text from DOCX
            paragraph_count: Number of paragraphs

        Returns:
            AI analysis result
        """
        # Truncate text if too long
        truncated_text = text[:15000]
        if len(text) > 15000:
            truncated_text += "\n\n[... content truncated ...]"

        prompt = f"""Analyze this Word document ({paragraph_count} paragraphs) and extract structured information.

Provide:

1. DOCUMENT TYPE: Identify what type of document this is (e.g., requirements doc, test scenarios, meeting notes, specification, user manual, etc.)

2. SUMMARY: Brief overview of the document content (2-3 sentences)

3. MAIN POINTS: List the key points, topics, or sections covered
   - Focus on the most important information
   - Include any headings or major sections

4. ACTION ITEMS: Extract any tasks, to-dos, or action items mentioned
   - Include assignees if mentioned
   - Note any deadlines or priorities

5. TEST SCENARIOS: If this is a test document, extract test cases/scenarios
   - Test case ID (if present)
   - Test description
   - Steps (if present)
   - Expected results (if present)
   - Priority (if mentioned)

6. IMPORTANT DETAILS: Any critical information like requirements, constraints, decisions, or dependencies

Format your response clearly with headers for each section.
Be specific and detailed for test scenarios if present.
Focus on information that would help a developer understand what needs to be implemented or tested."""

        return await self.ai_service.extract_text_content(
            text=truncated_text,
            prompt=prompt,
            max_tokens=3000
        )

    def _parse_analysis_response(self, content: str) -> Dict[str, Any]:
        """
        Parse AI analysis response into structured format.

        Args:
            content: Raw AI response

        Returns:
            Structured extracted content
        """
        result = {
            "document_type": "",
            "summary": "",
            "main_points": [],
            "action_items": [],
            "test_scenarios": []
        }

        try:
            lines = content.split('\n')
            current_section = None

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                line_upper = line.upper()

                # Detect section headers
                if 'DOCUMENT TYPE' in line_upper or line_upper.startswith('1.'):
                    current_section = 'document_type'
                    continue
                elif 'SUMMARY' in line_upper or line_upper.startswith('2.'):
                    current_section = 'summary'
                    continue
                elif 'MAIN POINTS' in line_upper or line_upper.startswith('3.'):
                    current_section = 'main_points'
                    continue
                elif 'ACTION ITEMS' in line_upper or line_upper.startswith('4.'):
                    current_section = 'action_items'
                    continue
                elif 'TEST SCENARIOS' in line_upper or line_upper.startswith('5.'):
                    current_section = 'test_scenarios'
                    continue
                elif 'IMPORTANT DETAILS' in line_upper or line_upper.startswith('6.'):
                    current_section = 'important_details'
                    continue

                # Add content to appropriate section
                if current_section == 'document_type':
                    if not result['document_type']:
                        result['document_type'] = line
                    else:
                        result['document_type'] += ' ' + line

                elif current_section == 'summary':
                    if result['summary']:
                        result['summary'] += ' ' + line
                    else:
                        result['summary'] = line

                elif current_section == 'main_points':
                    clean_line = line.lstrip('-*•›→ 0123456789.)')
                    if clean_line:
                        result['main_points'].append(clean_line)

                elif current_section == 'action_items':
                    clean_line = line.lstrip('-*•›→ 0123456789.)')
                    if clean_line:
                        result['action_items'].append(clean_line)

                elif current_section == 'test_scenarios':
                    clean_line = line.lstrip('-*•›→ 0123456789.)')
                    if clean_line:
                        # Try to parse as structured test scenario
                        if 'TC-' in clean_line or 'Test' in clean_line:
                            result['test_scenarios'].append({
                                "description": clean_line
                            })
                        elif result['test_scenarios']:
                            # Append to last test scenario
                            last_scenario = result['test_scenarios'][-1]
                            if 'details' not in last_scenario:
                                last_scenario['details'] = []
                            last_scenario['details'].append(clean_line)

            # If parsing failed, put everything in summary
            if not result['document_type'] and not result['summary'] and not any([
                result['main_points'],
                result['action_items'],
                result['test_scenarios']
            ]):
                result['summary'] = content

        except Exception as e:
            logger.warning(f"Error parsing DOCX analysis: {e}")
            result['summary'] = content

        return result
